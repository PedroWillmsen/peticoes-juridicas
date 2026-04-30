import os
import re

from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

from constants import (
    NORONHA_RODAPE,
    PARCERIA_RODAPE,
)

LOGO_PAR_ESQ = "logo_parceria_esquerda.png"
LOGO_PAR_DIR = "logo_parceria_direita.jpg"
LOGO_NORONHA = "logo_noronha.png"
ASS_ELEANDRO = "assinatura_eleandro.png"
ASS_MARILIA  = "assinatura_marilia.png"

_FONTE_NOME = "Segoe UI"
_FONTE_TAM  = 12
_MODELO     = "Parceria Marília + Eleandro"

_DATA_LINE_RE = re.compile(
    r'^(Titular|CNPJ|CPF|Banco|Agência|Agencia|Conta Corrente|Conta|'
    r'Telefone|E-mail|Email|Fax|Procuradores|Reclamante|Reclamado)[\s:/]',
    re.IGNORECASE,
)
_LABEL_ONLY_RE = re.compile(r'^[A-ZÀ-Úa-zà-ú][\w\s/()-]*:$')

_PROCESSO_RE = re.compile(
    r'^(Processo|CumSen|Execução|Cumprimento de Sentença|AP|Reclamação)\s*n[ºo°]?\s*',
    re.IGNORECASE,
)


def _is_data_line(linha: str) -> bool:
    s = linha.strip()
    if not s:
        return False
    if _DATA_LINE_RE.match(s):
        return True
    if _LABEL_ONLY_RE.match(s) and len(s) < 50:
        return True
    return False


def _set_font(run, name=None, size=None, bold=False):
    name = name or _FONTE_NOME
    size = size or _FONTE_TAM
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _parse_bold_runs(paragraph, text: str, size=None):
    size = size or _FONTE_TAM
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            _set_font(r, size=size, bold=True)
        elif part:
            r = paragraph.add_run(part)
            _set_font(r, size=size)


def _set_doc_style(doc):
    style = doc.styles["Normal"]
    style.font.name = _FONTE_NOME
    style.font.size = Pt(_FONTE_TAM)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), _FONTE_NOME)


def _configure_page(doc):
    sec = doc.sections[0]
    sec.page_width    = Mm(210)   # ← A4 largura
    sec.page_height   = Mm(297)   # ← A4 altura
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)
    sec.different_first_page_header_footer = True
    return sec


def _normalize_lines(texto: str):
    return [linha.rstrip() for linha in texto.split("\n")]


def _strip_trailing_block(texto: str, modelo: str):
    linhas = _normalize_lines(texto)

    if modelo == "Parceria Marília + Eleandro":
        blocos_remover = {
            "Eleandro Soares", "OAB/RS nº 70.936",
            "Marília Chemello Faviero", "OAB/RS 52.535",
            "Av. Ipiranga, nº 40, Cond. Trend City Center Offices, salas 1211 e 1911",
            "Bairro Praia de Belas – Porto Alegre – RS – CEP 90160-090",
            "eleandrosoares.adv@gmail.com | contato@mariliafaviero.adv.br",
            "Av. Ipiranga, nº 40, Cond. Trend City Center Offices, sala 1211",
            "eleandrosoares.adv@gmail.com",
            "Av. Senador Tarso Dutra, 565/1501 | Trend Offices",
            "CEP: 90690-140 | Porto Alegre | RS",
            "contato@mariliafaviero.adv.br",
        }
    else:
        blocos_remover = {
            "Marília Chemello Faviero Willmsen", "OAB/RS 52.535",
            "Geovana da Silva Freitas", "OAB/RS nº 59.771",
            "Ivandro Noronha de Freitas", "OAB/RS nº 97.120",
            "Eleandro Soares", "OAB/RS nº 70.936",
            "Rua dos Andradas, nº 1234, Ed. Santa Cruz, sala 1310",
            "Bairro Centro – Porto Alegre – RS – CEP 90020-008",
            "Fone: 51.3062-8510", "noronhasoaresadv@gmail.com",
        }

    resultado = []
    for linha in linhas:
        if linha.strip() in blocos_remover:
            continue
        resultado.append(linha)

    while resultado and resultado[-1].strip() == "":
        resultado.pop()

    return resultado


def _split_sections(linhas):
    cabecalho  = []
    corpo      = []
    fechamento = []
    state      = "header"

    for linha in linhas:
        if state == "header":
            cabecalho.append(linha)
            if _PROCESSO_RE.match(linha):
                state = "body"
            continue
        if state == "body":
            stripped = linha.strip()
            if stripped.startswith("Termos em que") or stripped.startswith("Nestes termos"):
                state = "closing"
                fechamento.append(linha)
            else:
                corpo.append(linha)
            continue
        fechamento.append(linha)

    return cabecalho, corpo, fechamento


def _add_header(section, modelo):
    header = section.first_page_header

    if modelo == "Noronha":
        p = header.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(LOGO_NORONHA):
            r = p.add_run()
            r.add_picture(LOGO_NORONHA, width=Cm(5.0))
        return

    table = header.add_table(rows=1, cols=2, width=Cm(16))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_default = header.paragraphs[0]
    p_default._element.addprevious(table._tbl)
    p_default.paragraph_format.space_before = Pt(0)
    p_default.paragraph_format.space_after  = Pt(0)
    p_default.paragraph_format.line_spacing = Pt(1)

    c1 = table.cell(0, 0)
    c2 = table.cell(0, 1)
    c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p1 = c1.paragraphs[0]
    p2 = c2.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if os.path.exists(LOGO_PAR_ESQ):
        r = p1.add_run()
        r.add_picture(LOGO_PAR_ESQ, width=Cm(3.6))
    if os.path.exists(LOGO_PAR_DIR):
        r = p2.add_run()
        r.add_picture(LOGO_PAR_DIR, width=Cm(4.5))


def _add_footer(section, modelo):
    footer = section.first_page_footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    texto = NORONHA_RODAPE if modelo == "Noronha" else PARCERIA_RODAPE
    r = p.add_run(texto)
    _set_font(r, size=7)


def _add_title(doc, linha):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(linha)
    _set_font(r, bold=True)


def _add_processo(doc, linha):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(linha)
    _set_font(r, bold=True)


def _add_body(doc, linha, is_intro=False):
    """
    is_intro=True  → parágrafo de abertura (reclamante/reclamado):
                     Noronha = bloco inteiro na régua 3 (left_indent)
                     Parceria = só 1ª linha na régua 3 (first_line_indent)
    is_intro=False → demais parágrafos: só 1ª linha na régua 3
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = 1.5

    if _MODELO == "Noronha" and is_intro:
        # Bloco inteiro na régua 3
        p.paragraph_format.left_indent       = Cm(3.0)
        p.paragraph_format.first_line_indent = Cm(0)
    else:
        # Só a primeira linha na régua 3
        p.paragraph_format.first_line_indent = Cm(3.0)
        p.paragraph_format.left_indent       = Cm(0)

    _parse_bold_runs(p, linha)


def _add_data_line(doc, linha, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent       = Cm(3.0)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(0)
    p.paragraph_format.line_spacing      = 1.5

    if ":" in linha:
        rotulo, _, valor = linha.partition(":")
        r1 = p.add_run(rotulo + ":")
        _set_font(r1, bold=True)        # "Titular:"  → negrito
        r2 = p.add_run(valor)
        _set_font(r2, bold=False)       # " ELEANDRO SOARES..." → normal
    else:
        r = p.add_run(linha)
        _set_font(r, bold=bold)


def _add_simple(doc, linha, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(0)
    p.paragraph_format.line_spacing      = 1.5
    r = p.add_run(linha)
    _set_font(r, bold=bold)


def _add_closing(doc, linha):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(3.0)
    p.paragraph_format.left_indent       = Cm(0)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(0)
    p.paragraph_format.line_spacing      = 1.5
    r = p.add_run(linha)
    _set_font(r)


def _blank(doc, after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)


def _add_signatures_parceria(doc):
    from docx.oxml import OxmlElement
    _blank(doc, 6)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Fix: força largura de 14cm + centraliza a tabela via XML
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(14 * 567)))  # 14cm em twips
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    left  = table.cell(0, 0)
    right = table.cell(0, 1)
    left.vertical_alignment  = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p1 = left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(ASS_ELEANDRO):
        p1.add_run().add_picture(ASS_ELEANDRO, width=Cm(3.3))
    else:
        r = p1.add_run("Eleandro Soares")
        _set_font(r, bold=True)
        p_oab = left.add_paragraph()
        p_oab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p_oab.add_run("OAB/RS nº 70.936"))

    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(ASS_MARILIA):
        p2.add_run().add_picture(ASS_MARILIA, width=Cm(3.5))
    else:
        r = p2.add_run("Marília Chemello Faviero")
        _set_font(r, bold=True)
        p_oab = right.add_paragraph()
        p_oab.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p_oab.add_run("OAB/RS 52.535"))

def _add_signatures_noronha(doc):
    _blank(doc, 6)
    nomes = [
        ("Marília Chemello Faviero",          "OAB/RS 52.535"),
        ("Geovana da Silva Freitas",          "OAB/RS nº 59.771"),
        ("Ivandro Noronha de Freitas",        "OAB/RS nº 97.120"),
        ("Eleandro Soares",                   "OAB/RS nº 70.936"),
    ]
    for nome, oab in nomes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(nome)
        _set_font(r, bold=True)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        r = p.add_run(oab)
        _set_font(r)


def gerar_docx(
    texto: str,
    modelo: str = "Parceria Marília + Eleandro",
    nome_arquivo: str = "peticao.docx",
):
    global _FONTE_NOME, _FONTE_TAM, _MODELO
    _MODELO     = modelo
    _FONTE_NOME = "Century Gothic" if modelo == "Noronha" else "Segoe UI"
    _FONTE_TAM  = 11               if modelo == "Noronha" else 12

    doc = Document()
    _set_doc_style(doc)
    section = _configure_page(doc)
    _add_header(section, modelo)
    _add_footer(section, modelo)

    linhas = _strip_trailing_block(texto, modelo)
    cabecalho, corpo, fechamento = _split_sections(linhas)

    for linha in cabecalho:
        if not linha.strip():
            _blank(doc)
        elif linha.startswith("AO JUÍZO") or linha.startswith("EXCELENTÍSSIMO"):
            _add_title(doc, linha)
        elif _PROCESSO_RE.match(linha):
            _blank(doc, 2)
            _add_processo(doc, linha)
        else:
            _add_simple(doc, linha)

    _blank(doc, 8)

    # FIX: rastreia o primeiro parágrafo do corpo (intro com reclamante/reclamado)
    primeiro_corpo = True
    for linha in corpo:
        if not linha.strip():
            _blank(doc)
        elif _is_data_line(linha):
            _add_data_line(doc, linha)
        else:
            _add_body(doc, linha, is_intro=primeiro_corpo)
            primeiro_corpo = False

    _blank(doc, 8)

    # FIX: fechamento sem linha em branco entre "Termos em que" e "Porto Alegre"
    for linha in fechamento:
        if linha.strip():
            _add_closing(doc, linha)

    if modelo == "Noronha":
        _add_signatures_noronha(doc)
    else:
        _add_signatures_parceria(doc)

    doc.save(nome_arquivo)
    return nome_arquivo